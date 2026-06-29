#!/usr/bin/env python3
"""
Hermes Skills Security Auditor — 审查 Hermes 技能的安全性

Usage:
    python audit_skills.py <skill_name>        # 审查指定技能
    python audit_skills.py --all               # 审查所有技能
    python audit_skills.py --output /tmp/r.json  # JSON 输出
    python audit_skills.py --report-doc /tmp/r.docx  # Word 报告

SAFETY: This tool performs READ-ONLY operations only. It reads skill
files but never writes or modifies them.
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path


# ═══════════════════════════════════════════════════
# PATTERNS: Things to look for in skill scripts
# ═══════════════════════════════════════════════════

# ── HIGH severity patterns ──
HIGH_PATTERNS = [
    ("shell=True", "Shell injection risk", "subprocess.run/call with shell=True allows command injection. Use shell=False with argument lists."),
    (r"\bos\.system\s*\(", "os.system() usage", "os.system() uses a shell and is injection-prone. Use subprocess.run with shell=False."),
    (r"\bos\.popen\s*\(", "os.popen() usage", "os.popen() is insecure. Use subprocess.Popen with proper arguments."),
    (r"\beval\s*\(", "eval() usage", "eval() can execute arbitrary code. Avoid using eval()."),
    (r"\bexec\s*\(", "exec() usage", "exec() can execute arbitrary code. Avoid using exec()."),
    (r"pickle\.loads?\s*\(", "Unsafe deserialization", "pickle.loads() can execute arbitrary code during deserialization. Use JSON or safer alternatives."),
    (r"rm\s+(-rf?\s+)?/", "Dangerous rm -rf /", "This command can destroy the entire filesystem."),
    (r"(?i)(api[_-]?key|secret|password|token|credential)\s*=\s*['\"][^'\"]{6,}['\"]", "Hardcoded credential", "Hardcoded API keys or passwords in scripts are security risks."),
    (r"chmod\s+-R\s+777", "Recursive world-writable permissions", "chmod -R 777 grants full access to everyone."),
    (r"subprocess\.(check_call|Popen|run)\(.*shell\s*=\s*True", "Shell=True in subprocess", "Subprocess with shell=True allows command injection."),
]

# ── MEDIUM severity patterns ──
MEDIUM_PATTERNS = [
    (r"chmod\s+777", "World-writable permissions", "Files with 777 permissions can be modified by any user."),
    (r"(?i)rm\s+-rf", "Force recursive delete", "rm -rf can silently delete large amounts of data. Validate paths."),
    (r"socket\.connect\s*\(", "Raw socket connection", "Raw sockets may bypass security controls."),
    (r"requests\.get|requests\.post|urllib\.request|curl\s+-o|wget\s+", "Network call to external URL", f"External network calls should be validated for security. Check URLs are expected."),
    (r"(?i)(password|passwd|secret)\s*=?\s*['\"][^'\"]+['\"]", "Potential credential in code", "Verify this is not a hardcoded credential."),
    (r"subprocess\.run\(.*shell\s*=\s*True", "Shell=True", "Consider using shell=False for safety."),
    (r"sudo\s+", "Sudo command usage", "Sudo commands may need user confirmation or have privilege implications."),
    (r"open\(['\"].*[\/\\]", "File write to path", "Check that file writes only target expected directories."),
    (r"\.write\s*\(", "Write operation", "Write operations should be reviewed to ensure they don't overwrite system files."),
    (r"os\.remove|os\.unlink|Path\(.*\)\.unlink", "File deletion", "File deletion operations should verify paths."),
    (r"shutil\.rmtree|shutil\.move|shutil\.copy", "File system modification", "File operations should validate source and destination paths."),
    (r"os\.chmod|os\.chown", "Permission modification", "Permission changes should be carefully scoped."),
    (r"service|systemctl|initctl", "Service management", "Service management commands can affect system availability."),
    (r"apt(?:-get)?\s+install|yum\s+install|pip\s+install|npm\s+install", "Package installation", "Package installation modifies system state."),
    (r"__import__\s*\(", "Dynamic import", "Dynamic imports with user input can be dangerous."),
    (r"compile\s*\(.*['\"]", "compile() with dynamic input", "Compiling code from strings can be unsafe."),
]

# ── LOW severity patterns ──
LOW_PATTERNS = [
    (r"#\s*TODO\s*:", "TODO comment", "TODOs may indicate incomplete security considerations."),
    (r"#\s*FIXME\s*:", "FIXME comment", "FIXME comments suggest known issues."),
    (r"#\s*HACK\s*:", "HACK comment", "HACK comments often bypass proper security patterns."),
    (r"print\s*\(.*password|print\s*\(.*secret|print\s*\(.*token", "Potential credential logging", "Avoid printing sensitive information to stdout."),
    (r"debug\s*=\s*True", "Debug mode enabled", "Debug mode may expose sensitive information."),
    (r"os\.environ\[", "Environment variable access", "Ensure environment variables don't contain secrets."),
    (r"tempfile|/tmp/", "Temporary file usage", "Temporary files in /tmp may be readable by other users."),
    (r"input\s*\(|sys\.stdin", "User input accepted", "User input should be validated and sanitized."),
]

# ── SKILL.md specific patterns ──
SKILLMD_PATTERNS = [
    ("sensitive_info", r"(?i)(password|api[_-]?key|secret|token|credential).*:.*['\"].{4,}['\"]", "Potential sensitive info in docs", "Sensitive information should not be in documentation."),
    ("dangerous_instruction", r"(?i)(rm\s+-rf\s+/|chmod\s+-R\s+777|> /dev/sda)", "Potentially dangerous instruction", "Instructions in skill docs that could be destructive."),
    ("missing_security", r"(?i)(security|safety|caution|warning|danger)", "Missing security section" if False else "", ""),  # Checked separately
]

# ── Whitelist: known safe network URLs ──
SAFE_URL_DOMAINS = [
    "github.com", "raw.githubusercontent.com", "pypi.org", "pypi.python.org",
    "files.pythonhosted.org", "npmjs.org", "registry.npmjs.org",
    "docker.io", "hub.docker.com", "quay.io",
    "api.openai.com", "api.anthropic.com",
    "googleapis.com", "graph.microsoft.com",
    "api.telegram.org", "api.slack.com",
    "localhost", "127.0.0.1",
]


def is_safe_url(url):
    """Check if a URL points to a known safe domain."""
    for domain in SAFE_URL_DOMAINS:
        if domain in url:
            return True
    return False


def scan_file_for_patterns(filepath, patterns, context=""):
    """Scan a file for security patterns and return findings."""
    findings = []
    try:
        content = filepath.read_text(encoding='utf-8', errors='replace')
    except (PermissionError, OSError, UnicodeDecodeError) as e:
        findings.append({
            "severity": "ERROR",
            "title": f"Cannot read file",
            "detail": f"{filepath.name}: {e}",
            "recommendation": "Check file permissions",
            "file": str(filepath),
            "line": 0,
        })
        return findings

    for i, line in enumerate(content.splitlines(), 1):
        stripped = line.strip()

        # ── Skip lines that are PATTERN DEFINITIONS (tuples of strings in pattern lists) ──
        # These look like:  ("pattern", "title", "recommendation")  or  (r"...", "...", "..."),
        if re.match(r'^\s*(\(".+",\s*".+",\s*".+"\)|\(r".+",\s*".+",\s*".+"\))\s*,?\s*$', stripped):
            continue

        for pattern, title, rec in patterns:
            # For network URL patterns, do additional SAME-LINE safety check
            if "Network call" in title:
                # Extract URL from the line
                urls = re.findall(r'https?://[^\s"\'\)]+', line)
                if urls:
                    all_safe = all(is_safe_url(u) for u in urls)
                    if all_safe:
                        continue  # Skip if URLs are safe

            if re.search(pattern, line):
                # Don't flag comments that are self-documenting
                if stripped.startswith('#'):
                    continue

                # Skip: line is inside a comment block (starts with #)
                # Skip: line is part of a pattern list entry (the pattern checking this IS the same line)
                if stripped.startswith("(") and stripped.rstrip().endswith("),") and pattern in stripped:
                    continue

                # ── Skip false positives inside comments about security ──
                if "shell=True" in pattern:
                    if re.search(r'shell\s*=\s*False\s*#', line):
                        continue
                    if '"shell=True"' in line or "'shell=True'" in line:
                        continue

                findings.append({
                    "severity": "HIGH" if pattern in [p[0] for p in HIGH_PATTERNS] else
                                ("MEDIUM" if any(p[0] == pattern for p in MEDIUM_PATTERNS) else "LOW"),
                    "title": title,
                    "detail": f"Line {i}: {stripped[:150]}",
                    "recommendation": rec,
                    "file": str(filepath),
                    "line": i,
                })

    return findings


def check_skillyml_security(filepath):
    """Check SKILL.md for security-relevant content."""
    findings = []
    try:
        content = filepath.read_text(encoding='utf-8')
    except (PermissionError, OSError):
        return findings

    lines = content.splitlines()

    # Check 1: Does it reference security/safety anywhere?
    has_security_section = bool(re.search(r'(?i)(security|safety|caution|warning|danger)', content))
    if not has_security_section:
        # Only flag if the skill has scripts that warrant it
        skill_dir = filepath.parent
        has_scripts = (skill_dir / "scripts").exists()
        if has_scripts:
            findings.append({
                "severity": "LOW",
                "title": "No security/safety section in SKILL.md",
                "detail": "Skills with scripts should document security considerations",
                "recommendation": "Add a Common Pitfalls or Safety section to SKILL.md",
                "file": str(filepath),
                "line": 0,
            })

    # Check 2: Look for sensitive info in docs
    for i, line in enumerate(lines, 1):
        m = re.search(r'(?i)(password|api[_-]?key|secret|credential)[:\s=]+["\']?([^"\'>\n]{8,})', line)
        if m and not line.strip().startswith("#"):
            findings.append({
                "severity": "HIGH",
                "title": "Potential credential in documentation",
                "detail": f"Line {i}: discloses what appears to be a credential",
                "recommendation": "Remove hardcoded credentials from documentation",
                "file": str(filepath),
                "line": i,
            })

    # Check 3: Check for destructive instructions
    for i, line in enumerate(lines, 1):
        if re.search(r'(?i)(rm\s+-rf\s+/|chmod\s+-R\s+777)', line) and not line.strip().startswith("#"):
            findings.append({
                "severity": "HIGH",
                "title": "Destructive command in documentation",
                "detail": f"Line {i}: {line.strip()[:150]}",
                "recommendation": "Warn users before suggesting destructive commands",
                "file": str(filepath),
                "line": i,
            })

    return findings


def audit_single_skill(skill_name):
    """Audit a single Hermes skill for security issues."""
    # Search for skill directory
    skill_paths = [
        Path.home() / ".hermes" / "skills",
        Path.home() / ".hermes" / "hermes-agent" / "skills" if (Path.home() / ".hermes" / "hermes-agent" / "skills").exists() else None,
    ]
    skill_dir = None

    for base in [p for p in skill_paths if p]:
        # Search recursively for a skill directory with this name
        for path in base.rglob(skill_name):
            if path.is_dir() and (path / "SKILL.md").exists():
                skill_dir = path
                break
        if skill_dir:
            break

    if not skill_dir:
        # Also try direct paths
        for base in [p for p in skill_paths if p]:
            direct = base / skill_name
            if (direct / "SKILL.md").exists():
                skill_dir = direct
                break
            # Check subdirectories
            for cat_dir in base.iterdir():
                if cat_dir.is_dir():
                    possible = cat_dir / skill_name
                    if (possible / "SKILL.md").exists():
                        skill_dir = possible
                        break
            if skill_dir:
                break

    if not skill_dir:
        return {
            "skill": skill_name,
            "status": "not_found",
            "findings": [],
            "summary": {},
        }

    findings = []

    # 1. Check SKILL.md
    skillyml = skill_dir / "SKILL.md"
    if skillyml.exists():
        findings.extend(check_skillyml_security(skillyml))

    # 2. Check scripts directory
    scripts_dir = skill_dir / "scripts"
    if scripts_dir.exists():
        # Also scan init files and subdirectories
        for ext in ("*.py", "*.sh", "*.js", "*.rb"):
            for script_file in scripts_dir.rglob(ext):
                if script_file.is_file() and script_file.stat().st_size < 500000:  # Skip very large files
                    findings.extend(scan_file_for_patterns(script_file, HIGH_PATTERNS))
                    findings.extend(scan_file_for_patterns(script_file, MEDIUM_PATTERNS))
                    findings.extend(scan_file_for_patterns(script_file, LOW_PATTERNS))

    # 3. Check references for sensitive info
    refs_dir = skill_dir / "references"
    if refs_dir.exists():
        for ref_file in refs_dir.rglob("*.md"):
            if ref_file.is_file():
                findings.extend(check_skillyml_security(ref_file))

    # Sort by severity
    severity_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "INFO": 3, "ERROR": 4}
    findings.sort(key=lambda x: severity_order.get(x.get("severity", "INFO"), 99))

    # Summary
    severity_counts = {}
    for f in findings:
        sev = f.get("severity", "INFO")
        severity_counts[sev] = severity_counts.get(sev, 0) + 1

    score = 100
    score -= severity_counts.get("HIGH", 0) * 20
    score -= severity_counts.get("MEDIUM", 0) * 8
    score -= severity_counts.get("LOW", 0) * 3
    score = max(0, min(100, score))

    return {
        "skill": skill_name,
        "path": str(skill_dir),
        "status": "audited",
        "summary": {
            "score": score,
            "total_findings": len(findings),
            "by_severity": severity_counts,
            "has_scripts": (scripts_dir.exists()),
        },
        "findings": findings,
    }


def audit_all_skills():
    """Audit all available Hermes skills."""
    skills_base = [
        Path.home() / ".hermes" / "skills",
    ]
    results = []

    for base in skills_base:
        if not base.exists():
            continue
        for category_dir in base.iterdir():
            if not category_dir.is_dir():
                continue
            for skill_dir in category_dir.iterdir():
                if not skill_dir.is_dir():
                    continue
                if not (skill_dir / "SKILL.md").exists():
                    continue
                result = audit_single_skill(skill_dir.name)
                if result["status"] == "audited":
                    results.append(result)

    return results


def generate_summary(results):
    """Generate human-readable summary."""
    if isinstance(results, dict) and "skill" in results:
        results = [results]

    lines = []
    overall_score = 0
    total_high = 0
    total_medium = 0

    for r in results:
        s = r["summary"]
        overall_score += s.get("score", 0)
        total_high += s.get("by_severity", {}).get("HIGH", 0)
        total_medium += s.get("by_severity", {}).get("MEDIUM", 0)

        lines.append(f"\n{'='*60}")
        lines.append(f"📦 Skill: {r['skill']}")
        lines.append(f"📁 Path: {r.get('path', 'N/A')}")
        lines.append(f"📊 Score: {s.get('score', 'N/A')}/100")
        sc = s.get("by_severity", {})
        parts = []
        for sev in ("HIGH", "MEDIUM", "LOW", "INFO"):
            if sev in sc:
                parts.append(f"{sev}: {sc[sev]}")
        lines.append(f"🔍 Findings: {', '.join(parts) if parts else 'None'}")

        if r.get("findings"):
            for f in r["findings"]:
                sev = f.get("severity", "INFO")
                icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢", "INFO": "ℹ️", "ERROR": "❌"}
                lines.append(f"  {icon.get(sev, '•')} [{sev}] {f['title']}")
                lines.append(f"     File: {Path(f.get('file', '')).relative_to(Path.home()) if f.get('file') else 'N/A'}:{f.get('line', 0)}")
                lines.append(f"     Detail: {f.get('detail', '')[:120]}")
                lines.append(f"     → {f.get('recommendation', '')}")
                lines.append("")

    if results:
        avg_score = overall_score / len(results)
        lines.append(f"\n{'='*60}")
        lines.append(f"OVERALL: {len(results)} skills audited")
        lines.append(f"Average Score: {avg_score:.0f}/100")
        lines.append(f"Total HIGH: {total_high} | Total MEDIUM: {total_medium}")
        lines.append(f"{'='*60}")

    return "\n".join(lines)


def generate_word_report(results, output_path):
    """Generate a Word document security report for skills audit."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None, "python-docx not installed"

    if isinstance(results, dict) and "skill" in results:
        results = [results]

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hermes Skills Security Audit Report")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{len(results)} skills audited")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Overall score
    total_score = sum(r["summary"]["score"] for r in results)
    avg_score = total_score / len(results) if results else 0

    total_high = sum(r["summary"]["by_severity"].get("HIGH", 0) for r in results)
    total_med = sum(r["summary"]["by_severity"].get("MEDIUM", 0) for r in results)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Average Security Score: {avg_score:.0f}/100")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1B, 0x8A, 0x3A) if avg_score >= 70 else RGBColor(0xCC, 0x88, 0x00)

    doc.add_page_break()

    for r_data in results:
        s = r_data["summary"]
        h = doc.add_heading(r_data["skill"], level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

        doc.add_paragraph(f"Score: {s['score']}/100 | Path: {r_data.get('path', 'N/A')}")
        sc = s.get("by_severity", {})
        for sev in ("HIGH", "MEDIUM", "LOW", "INFO"):
            if sev in sc:
                p = doc.add_paragraph()
                r = p.add_run(f"  {sev}: {sc[sev]}")
                r.bold = True
                color_map = {"HIGH": RGBColor(0xCC, 0x33, 0x33),
                             "MEDIUM": RGBColor(0xCC, 0x88, 0x00),
                             "LOW": RGBColor(0x66, 0x99, 0x33),
                             "INFO": RGBColor(0x66, 0x66, 0x66)}
                r.font.color.rgb = color_map.get(sev)

        if r_data.get("findings"):
            for f in r_data["findings"]:
                p = doc.add_paragraph()
                r = p.add_run(f"[{f['severity']}] {f['title']}")
                r.bold = True
                r.font.size = Pt(10)
                fp = Path(f.get('file', ''))
                rel = fp.relative_to(Path.home()) if fp.is_absolute() else fp
                doc.add_paragraph(f"  File: ~/{rel}:{f.get('line', 0)}")
                doc.add_paragraph(f"  {f.get('detail', '')[:200]}")
                p_rec = doc.add_paragraph(f"  → {f.get('recommendation', '')}")
                p_rec.runs[0].font.color.rgb = RGBColor(0x2E, 0x6B, 0xA4)
                doc.add_paragraph()

        doc.add_page_break()

    doc.save(output_path)
    return output_path, None


def main():
    parser = argparse.ArgumentParser(
        description="Hermes Skills Security Auditor — Review skills for security issues",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
SAFETY: READ-ONLY — this tool scans skill files but never writes or modifies them.
        """
    )
    parser.add_argument("skill", nargs="?", help="Skill name to audit (omit or use --all for all skills)")
    parser.add_argument("--all", "-a", action="store_true", help="Audit all skills")
    parser.add_argument("--output", "-o", help="Save results as JSON")
    parser.add_argument("--report-doc", help="Generate Word audit report")
    parser.add_argument("--text", action="store_true", help="Print readable report")

    args = parser.parse_args()

    if args.all or not args.skill:
        results = audit_all_skills()
    else:
        result = audit_single_skill(args.skill)
        results = [result]

    # Filter out not-found
    results = [r for r in results if r["status"] == "audited"]

    if not results:
        print("No skills found to audit.")
        sys.exit(1)

    # Output
    if args.output:
        with open(args.output, "w") as f:
            json.dump(results if len(results) > 1 else results[0], f, indent=2, ensure_ascii=False)
        print(f"✅ Results saved to: {args.output}")

    if args.report_doc:
        path, err = generate_word_report(results, args.report_doc)
        if path:
            print(f"✅ Word report: {path}")
        else:
            print(f"❌ Word report failed: {err}")

    if args.text or not (args.output or args.report_doc):
        print(generate_summary(results))

    # Summary line
    total_high = sum(r["summary"]["by_severity"].get("HIGH", 0) for r in results)
    total_med = sum(r["summary"]["by_severity"].get("MEDIUM", 0) for r in results)
    avg_score = sum(r["summary"]["score"] for r in results) / len(results)
    print(f"\n{'='*50}")
    print(f"Skills: {len(results)} | Avg Score: {avg_score:.0f}/100 | HIGH: {total_high} | MEDIUM: {total_med}")
    print(f"{'='*50}")


if __name__ == "__main__":
    main()
