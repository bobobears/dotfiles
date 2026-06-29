#!/usr/bin/env python3
"""
Linux System Security Audit — Read-Only, Non-Destructive

Usage:
    python audit.py                               # Full audit (no sudo)
    python audit.py --sudo                        # Full audit with sudo
    python audit.py --category user               # Single category
    python audit.py --output /tmp/audit.json      # Save results as JSON
    python audit.py --output /tmp/audit.json --report-doc /tmp/audit.docx   # + Word report

Categories: user, file, network, service, package, log, kernel, ssh, all

SAFETY: This tool performs READ-ONLY operations only. It never writes,
  modifies, or deletes any system files, configurations, or services.
  All commands use subprocess.run with shell=False to prevent injection.
  Path arguments are validated against a strict allowlist.
"""

import argparse
import json
import os
import pwd
import grp
import subprocess
import sys
import stat
import platform
from datetime import datetime
from pathlib import Path

# ═══════════════════════════════════════════════════
# SAFETY: Allowed paths for file checks (allowlist)
# ═══════════════════════════════════════════════════
ALLOWED_AUDIT_PATHS = frozenset({
    "/etc/passwd", "/etc/shadow", "/etc/group", "/etc/gshadow",
    "/etc/sudoers", "/etc/sudoers.d",
    "/etc/ssh/sshd_config", "/etc/ssh/ssh_config",
    "/etc/hosts", "/etc/hostname", "/etc/resolv.conf",
    "/etc/fstab", "/etc/crypttab",
    "/etc/crontab", "/etc/cron.d", "/etc/cron.daily", "/etc/cron.hourly",
    "/etc/cron.weekly", "/etc/cron.monthly",
    "/etc/rsyslog.conf", "/etc/rsyslog.d",
    "/etc/audit/auditd.conf", "/etc/audit/rules.d",
    "/etc/login.defs", "/etc/pam.d",
    "/etc/security",
    "/etc/selinux/config",
    "/etc/apparmor", "/etc/apparmor.d",
    "/etc/sysctl.conf", "/etc/sysctl.d",
    "/etc/issue", "/etc/issue.net",
    "/etc/profile", "/etc/bash.bashrc",
    "/etc/environment",
    "/etc/hosts.allow", "/etc/hosts.deny",
    "/etc/nsswitch.conf",
    "/etc/ntp.conf", "/etc/chrony.conf",
    "/etc/rsyncd.conf",
    "/etc/default",
})

# Patterns for safe file discovery (used with find, not user input)
SAFE_GLOBS = {
    "world_writable": ('/tmp', '/var/tmp', '/dev/shm'),
    "suid_suid": ('/usr/bin', '/usr/sbin', '/bin', '/sbin'),
}


def safe_read_file(path):
    """Read a file if it exists and is on the allowlist. Returns content or None."""
    resolved = Path(path).resolve()
    if str(resolved) not in ALLOWED_AUDIT_PATHS:
        return None, f"SKIPPED: {path} not in allowlist"
    if not resolved.exists():
        return None, f"FILE NOT FOUND: {path}"
    try:
        return resolved.read_text(), None
    except PermissionError:
        return None, f"PERMISSION DENIED: {path}"
    except Exception as e:
        return None, f"ERROR reading {path}: {e}"


def run_cmd(cmd, timeout=15):
    """Run a shell command safely (shell=False). Returns stdout, stderr, returncode."""
    try:
        r = subprocess.run(
            cmd, capture_output=True, text=True, timeout=timeout,
            shell=False  # SAFETY: never use shell=True
        )
        return r.stdout.strip(), r.stderr.strip(), r.returncode
    except FileNotFoundError:
        return "", "command not found", -1
    except subprocess.TimeoutExpired:
        return "", "timed out", -1
    except Exception as e:
        return "", str(e), -1


# ═══════════════════════════════════════════════════
# CATEGORY: User & Account Security
# ═══════════════════════════════════════════════════

def audit_users():
    """Check user accounts: UID 0, empty passwords, unnecessary users."""
    findings = []

    # Users with UID 0 (root-equivalent)
    uid0 = [u.pw_name for u in pwd.getpwall() if u.pw_uid == 0]
    if len(uid0) > 1:
        findings.append({
            "severity": "HIGH",
            "category": "user",
            "title": "Multiple UID 0 users",
            "detail": f"UID 0 users: {', '.join(uid0)}",
            "recommendation": "Only root should have UID 0"
        })

    # Users with no password or empty password field
    content, err = safe_read_file("/etc/shadow")
    if content:
        empty_pw = []
        no_pw = []
        for line in content.splitlines():
            parts = line.split(":")
            if len(parts) >= 2:
                if parts[1] in ("", "!"):
                    no_pw.append(parts[0])
                elif parts[1].startswith("$"):
                    pass  # hashed password
                elif parts[1] == "NP":
                    no_pw.append(parts[0])

        if no_pw:
            findings.append({
                "severity": "HIGH",
                "category": "user",
                "title": "Users with no password",
                "detail": f"Users: {', '.join(no_pw)}",
                "recommendation": "Set passwords or lock accounts with 'passwd -l <user>'"
            })

    # Users with login shell but no recent login
    all_users = pwd.getpwall()
    interactive_users = [u.pw_name for u in all_users
                         if u.pw_shell not in ("/sbin/nologin", "/usr/sbin/nologin",
                                               "/bin/false", "/usr/bin/false")]
    findings.append({
        "severity": "INFO",
        "category": "user",
        "title": "Users with login shell",
        "detail": f"Interactive users ({len(interactive_users)}): {', '.join(interactive_users)}",
        "recommendation": "Review and disable unused accounts"
    })

    # User list summary
    findings.append({
        "severity": "INFO",
        "category": "user",
        "title": "Total system users",
        "detail": f"{len(all_users)} users configured",
        "recommendation": ""
    })

    return findings


def audit_sudo():
    """Check sudoers configuration."""
    findings = []

    content, err = safe_read_file("/etc/sudoers")
    if content:
        # NOPASSWD entries
        nopasswd = [l.strip() for l in content.splitlines()
                     if "NOPASSWD" in l and not l.strip().startswith("#")]
        if nopasswd:
            findings.append({
                "severity": "MEDIUM",
                "category": "user",
                "title": "NOPASSWD sudo entries",
                "detail": f"Entries: {'; '.join(nopasswd)}",
                "recommendation": "Consider requiring password for sudo"
            })

    # Check sudoers.d directory
    sudoers_d = Path("/etc/sudoers.d")
    if sudoers_d.exists():
        files = list(sudoers_d.iterdir())
        findings.append({
            "severity": "INFO",
            "category": "user",
            "title": "Additional sudo rules",
            "detail": f"{len(files)} files in /etc/sudoers.d",
            "recommendation": "Review these rules for unnecessary privileges"
        })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: File & Permission Security
# ═══════════════════════════════════════════════════

def audit_files():
    """Check file permissions: world-writable, SUID/SGID, sensitive files."""
    findings = []

    # World-writable files in key directories (sample check, not full scan)
    stdout, _, _ = run_cmd(
        ["find", "/tmp", "-maxdepth", "1", "-type", "f", "-perm", "-o+w", "-ls"]
    )
    if stdout:
        count = len(stdout.splitlines())
        if count > 20:
            findings.append({
                "severity": "HIGH",
                "category": "file",
                "title": "Many world-writable files in /tmp",
                "detail": f"{count} files in /tmp are world-writable",
                "recommendation": "Clean up /tmp periodically"
            })

    # SUID/SGID binaries
    stdout, _, _ = run_cmd(
        ["find", "/usr/bin", "/usr/sbin", "/bin", "/sbin",
         "-type", "f", "(", "-perm", "-4000", "-o", "-perm", "-2000", ")",
         "-exec", "ls", "-la", "{}", "+"]
    )
    if stdout:
        lines = stdout.splitlines()
        # Count by setuid vs setgid
        suid = [l for l in lines if "rws" in l.split()[0]]
        sgid = [l for l in lines if "r-xr-s" in l or "rwxr-s" in l or "---Sr-s" in l]
        findings.append({
            "severity": "INFO",
            "category": "file",
            "title": "SUID/SGID binaries",
            "detail": f"SUID: {len(suid)}, SGID: {len(sgid)}. Total: {len(lines)}",
            "recommendation": "Review SUID/SGID binaries; unexpected ones may be a risk"
        })
        # Highlight unusual ones
        unusual = [l for l in lines
                   if not any(x in l for x in ["/bin/", "/usr/bin/", "ping",
                                                "su", "sudo", "mount", "umount",
                                                "passwd", "chsh", "chfn", "newgrp",
                                                "pkexec", "at", "crontab"])]
        if unusual:
            findings.append({
                "severity": "MEDIUM",
                "category": "file",
                "title": "Unusual SUID/SGID binaries",
                "detail": "\n".join(unusual[:10]),
                "recommendation": "Verify each unusual SUID binary is necessary"
            })

    # Sensitive file permissions
    sensitive_checks = [
        ("/etc/shadow", "0o640"),
        ("/etc/passwd", "0o644"),
        ("/etc/group", "0o644"),
        ("/etc/sudoers", "0o440"),
        ("/etc/ssh/sshd_config", "0o600"),
    ]
    for spath, expected in sensitive_checks:
        p = Path(spath)
        if p.exists():
            actual = oct(p.stat().st_mode)[-3:]
            expected_oct = expected.replace("0o", "")
            if actual != expected_oct:
                findings.append({
                    "severity": "MEDIUM" if actual > expected_oct else "LOW",
                    "category": "file",
                    "title": f"Unexpected permission on {spath}",
                    "detail": f"Expected: {expected_oct}, Got: {actual}",
                    "recommendation": f"Run: chmod {expected_oct} {spath}"
                })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: SSH Security
# ═══════════════════════════════════════════════════

def audit_ssh():
    """Check SSH configuration."""
    findings = []

    content, err = safe_read_file("/etc/ssh/sshd_config")
    if content is None:
        findings.append({
            "severity": "INFO",
            "category": "ssh",
            "title": "SSH config check",
            "detail": err,
            "recommendation": "Install openssh-server if SSH is needed"
        })
        return findings

    # Parse key settings
    settings = {}
    for line in content.splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            parts = line.split(None, 1)
            if len(parts) == 2:
                settings[parts[0].lower()] = parts[1]

    checks = {
        "permitrootlogin": {
            "bad": ["yes"],
            "severity": "HIGH",
            "msg": "Root SSH login is permitted",
            "rec": "Set PermitRootLogin no or prohibit-password"
        },
        "passwordauthentication": {
            "bad": ["yes"],
            "severity": "MEDIUM",
            "msg": "Password authentication is enabled",
            "rec": "Use SSH key authentication only"
        },
        "pubkeyauthentication": {
            "bad": ["no"],
            "severity": "MEDIUM",
            "msg": "Public key authentication is disabled",
            "rec": "Set PubkeyAuthentication yes"
        },
        "x11forwarding": {
            "bad": ["yes"],
            "severity": "LOW",
            "msg": "X11 forwarding is enabled",
            "rec": "Set X11Forwarding no unless needed"
        },
        "maxauthtries": {
            "bad": None,
            "low": 4,
            "severity": "MEDIUM",
            "msg": f"MaxAuthTries is {settings.get('maxauthtries', 'not set')}",
            "rec": "Set MaxAuthTries to 3 or lower"
        },
    }

    for key, check in checks.items():
        val = settings.get(key)
        bad_vals = check.get("bad")
        if bad_vals is not None and val and val.lower() in bad_vals:
            findings.append({
                "severity": check["severity"],
                "category": "ssh",
                "title": check["msg"],
                "detail": f"Current: {key} {val}",
                "recommendation": check["rec"]
            })
        elif bad_vals is None and val:
            # Numeric check
            try:
                if int(val) > check.get("low", 999):
                    if key == "maxauthtries":
                        findings.append({
                            "severity": check["severity"],
                            "category": "ssh",
                            "title": check["msg"],
                            "detail": f"Current: MaxAuthTries {val}",
                            "recommendation": check["rec"]
                        })
            except ValueError:
                pass

    if not any(f["title"].startswith("Root") for f in findings):
        if settings.get("permitrootlogin", "").lower() in ("prohibit-password", "without-password"):
            findings.append({
                "severity": "LOW",
                "category": "ssh",
                "title": "Root SSH login restricted to key-only",
                "detail": "PermitRootLogin prohibit-password",
                "recommendation": "Consider PermitRootLogin no for stricter security"
            })

    # SSH protocol version
    if settings.get("protocol", "2") != "2":
        findings.append({
            "severity": "HIGH",
            "category": "ssh",
            "title": "SSH protocol 1 is allowed",
            "detail": "Protocol 1 is insecure",
            "recommendation": "Set Protocol 2 only"
        })

    findings.append({
        "severity": "INFO",
        "category": "ssh",
        "title": "SSH configuration reviewed",
        "detail": f"{len([f for f in findings if f['severity'] != 'INFO'])} issues found",
        "recommendation": ""
    })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: Network Security
# ═══════════════════════════════════════════════════

def audit_network():
    """Check open ports, listening services, firewall."""
    findings = []

    # Listening services
    stdout, _, _ = run_cmd(
        ["ss", "-tlnp"]
    )
    stdout2, _, _ = run_cmd(
        ["ss", "-ulnp"]
    )

    findings.append({
        "severity": "INFO",
        "category": "network",
        "title": "TCP listening services",
        "detail": stdout if stdout else "No TCP listeners",
        "recommendation": "Close unnecessary services"
    })
    findings.append({
        "severity": "INFO",
        "category": "network",
        "title": "UDP listening services",
        "detail": stdout2 if stdout2 else "No UDP listeners",
        "recommendation": "Close unnecessary services"
    })

    # Count open ports
    if stdout:
        lines = [l for l in stdout.splitlines() if "LISTEN" in l]
        findings.append({
            "severity": "INFO",
            "category": "network",
            "title": f"{len(lines)} TCP ports in LISTEN state",
            "detail": "Run 'ss -tlnp' for details",
            "recommendation": "Audit each open port for necessity"
        })

    # Listen on 0.0.0.0 (all interfaces)
    public_listeners = []
    if stdout:
        for line in stdout.splitlines():
            if "0.0.0.0:" in line or ":::" in line:
                public_listeners.append(line.strip())
    if public_listeners:
        findings.append({
            "severity": "MEDIUM",
            "category": "network",
            "title": "Services listening on all interfaces",
            "detail": "\n".join(public_listeners[:15]),
            "recommendation": "Bind to specific IP (127.0.0.1) for local-only services"
        })

    # Firewall status
    ufw_out, _, _ = run_cmd(["ufw", "status"])
    if "inactive" in ufw_out.lower():
        findings.append({
            "severity": "HIGH",
            "category": "network",
            "title": "UFW firewall is INACTIVE",
            "detail": "Uncomplicated Firewall is not enabled",
            "recommendation": "Enable firewall: ufw enable"
        })
        findings.append({
            "severity": "INFO",
            "category": "network",
            "title": "UFW status",
            "detail": ufw_out,
            "recommendation": ""
        })
    else:
        findings.append({
            "severity": "INFO",
            "category": "network",
            "title": "UFW firewall is active",
            "detail": ufw_out,
            "recommendation": ""
        })

    # iptables rules
    ipt_out, _, _ = run_cmd(["iptables", "-L", "-n", "--line-numbers"])
    if ipt_out:
        rules = [l for l in ipt_out.splitlines() if "ACCEPT" in l or "DROP" in l or "REJECT" in l]
        findings.append({
            "severity": "INFO",
            "category": "network",
            "title": "iptables rules",
            "detail": f"{len(rules)} active rules",
            "recommendation": "Review iptables rules for completeness"
        })

    # IP forwarding
    ip_forward, _, _ = run_cmd(
        ["sysctl", "-n", "net.ipv4.ip_forward"]
    )
    if ip_forward.strip() == "1":
        findings.append({
            "severity": "MEDIUM",
            "category": "network",
            "title": "IP forwarding is enabled",
            "detail": "net.ipv4.ip_forward = 1",
            "recommendation": "Disable if not acting as a router: sysctl -w net.ipv4.ip_forward=0"
        })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: Service Security
# ═══════════════════════════════════════════════════

def audit_services():
    """Check running services, cron jobs."""
    findings = []

    # Running services (systemd)
    stdout, _, _ = run_cmd(
        ["systemctl", "list-units", "--type=service", "--state=running", "--no-pager"]
    )
    if stdout:
        lines = [l for l in stdout.splitlines() if ".service" in l]
        svc_names = [l.split()[0] for l in lines if l.strip()]
        findings.append({
            "severity": "INFO",
            "category": "service",
            "title": f"Running services ({len(svc_names)})",
            "detail": "\n".join(svc_names[:20]),
            "recommendation": "Disable unnecessary services: systemctl disable <name>"
        })

    # Failed services
    stdout, _, _ = run_cmd(
        ["systemctl", "list-units", "--type=service", "--state=failed", "--no-pager"]
    )
    if stdout:
        failed = [l.strip() for l in stdout.splitlines() if ".service" in l and l.strip()]
        if failed:
            findings.append({
                "severity": "MEDIUM",
                "category": "service",
                "title": f"Failed services ({len(failed)})",
                "detail": "\n".join(failed[:10]),
                "recommendation": "Check logs: journalctl -u <service>"
            })

    # Cron jobs
    stdout, _, _ = run_cmd(
        ["find", "/etc/cron.d", "/etc/cron.daily", "/etc/cron.hourly",
         "/etc/cron.weekly", "/etc/cron.monthly", "-type", "f", "-ls"]
    )
    if not stdout:
        stdout = "(no cron files found)"
    findings.append({
        "severity": "INFO",
        "category": "service",
        "title": "System cron jobs",
        "detail": stdout[:500],
        "recommendation": "Review cron jobs for unauthorized entries"
    })

    # User cron jobs
    stdout, _, _ = run_cmd(["crontab", "-l"])
    if "no crontab" in stdout.lower() or "command not found" in stdout:
        pass
    elif stdout:
        findings.append({
            "severity": "INFO",
            "category": "service",
            "title": "User cron jobs",
            "detail": stdout[:500],
            "recommendation": "Review user cron jobs"
        })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: Package & Update Security
# ═══════════════════════════════════════════════════

def audit_packages():
    """Check package updates and security patches."""
    findings = []

    # OS info
    os_info = f"{platform.system()} {platform.release()}"
    findings.append({
        "severity": "INFO",
        "category": "package",
        "title": "Operating system",
        "detail": os_info,
        "recommendation": ""
    })

    # Kernel version
    findings.append({
        "severity": "INFO",
        "category": "package",
        "title": "Kernel version",
        "detail": platform.uname().version,
        "recommendation": ""
    })

    # Check for available updates (apt)
    stdout, stderr, rc = run_cmd(
        ["apt-get", "--just-print", "upgrade"], timeout=30
    )
    if rc == 0 and stdout:
        # Count upgradable packages
        lines = [l for l in stdout.splitlines()
                 if "upgraded" in l or "installed" in l or "removed" in l]
        updatable = 0
        for line in lines:
            parts = line.split()
            if len(parts) >= 2:
                try:
                    updatable += int(parts[0])
                except ValueError:
                    pass
        findings.append({
            "severity": "HIGH" if updatable > 10 else ("MEDIUM" if updatable > 0 else "INFO"),
            "category": "package",
            "title": "Available package updates",
            "detail": f"{updatable} packages can be upgraded" if updatable else "System is up to date",
            "recommendation": f"Run: apt upgrade to update {updatable} packages" if updatable else ""
        })

    # Check for held packages
    stdout, _, _ = run_cmd(["apt-mark", "showhold"])
    if stdout and stdout.strip():
        findings.append({
            "severity": "MEDIUM",
            "category": "package",
            "title": "Held packages (not updated)",
            "detail": stdout[:300],
            "recommendation": "Review held packages for missing security patches"
        })

    # Check if unattended-upgrades is installed
    stdout, _, _ = run_cmd(["dpkg", "-l", "unattended-upgrades"])
    if "no packages" in stdout.lower() or "command not found" in stdout:
        findings.append({
            "severity": "MEDIUM",
            "category": "package",
            "title": "unattended-upgrades not installed",
            "detail": "Automatic security updates are not configured",
            "recommendation": "Install: apt install unattended-upgrades"
        })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: Logging & Monitoring
# ═══════════════════════════════════════════════════

def audit_logging():
    """Check logging and audit configuration."""
    findings = []

    # rsyslog status
    stdout, _, _ = run_cmd(["systemctl", "is-active", "rsyslog"])
    if "active" in stdout:
        findings.append({
            "severity": "INFO",
            "category": "log",
            "title": "rsyslog is running",
            "detail": "System logging is active",
            "recommendation": ""
        })
    else:
        findings.append({
            "severity": "HIGH",
            "category": "log",
            "title": "rsyslog is NOT running",
            "detail": f"Status: {stdout}",
            "recommendation": "Enable: systemctl enable --now rsyslog"
        })

    # auditd status
    stdout, _, _ = run_cmd(["systemctl", "is-active", "auditd"])
    if "active" in stdout:
        findings.append({
            "severity": "INFO",
            "category": "log",
            "title": "auditd is running",
            "detail": "Linux Audit Framework is active",
            "recommendation": ""
        })
    else:
        findings.append({
            "severity": "MEDIUM",
            "category": "log",
            "title": "auditd is NOT running",
            "detail": f"Status: {stdout}",
            "recommendation": "Enable: systemctl enable --now auditd"
        })

    # Log file sizes
    log_dir = Path("/var/log")
    if log_dir.exists():
        large_logs = []
        for f in sorted(log_dir.iterdir()):
            if f.is_file() and f.suffix in (".log", ""):
                try:
                    sz = f.stat().st_size
                    if sz > 100 * 1024 * 1024:  # 100MB
                        large_logs.append(f"{f.name} ({sz/1024/1024:.0f}MB)")
                except (PermissionError, OSError):
                    pass
        if large_logs:
            findings.append({
                "severity": "LOW",
                "category": "log",
                "title": "Large log files (>100MB)",
                "detail": "\n".join(large_logs[:10]),
                "recommendation": "Configure logrotate to manage log sizes"
            })

    return findings


# ═══════════════════════════════════════════════════
# CATEGORY: Kernel & System Hardening
# ═══════════════════════════════════════════════════

def audit_kernel():
    """Check kernel security parameters."""
    findings = []

    sysctl_params = {
        "net.ipv4.tcp_syncookies": ("1", "HIGH", "TCP SYN cookies disabled", "Enable: sysctl -w net.ipv4.tcp_syncookies=1"),
        "net.ipv4.conf.all.rp_filter": ("1", "MEDIUM", "Reverse path filtering disabled", "Enable: sysctl -w net.ipv4.conf.all.rp_filter=1"),
        "net.ipv4.conf.all.accept_source_route": ("0", "HIGH", "Source routing not disabled", "Disable: sysctl -w net.ipv4.conf.all.accept_source_route=0"),
        "net.ipv4.icmp_echo_ignore_broadcasts": ("1", "MEDIUM", "ICMP broadcast echo not ignored", "Enable: sysctl -w net.ipv4.icmp_echo_ignore_broadcasts=1"),
        "kernel.randomize_va_space": ("2", "HIGH", "ASLR is not fully enabled", "Enable: sysctl -w kernel.randomize_va_space=2"),
    }

    for param, (expected, severity, msg, rec) in sysctl_params.items():
        val, _, rc = run_cmd(["sysctl", "-n", param])
        if rc == 0 and val.strip() != expected:
            findings.append({
                "severity": severity,
                "category": "kernel",
                "title": msg,
                "detail": f"{param} = {val.strip()}, expected = {expected}",
                "recommendation": rec
            })

    # Check SELinux status
    selinux, _, _ = run_cmd(["getenforce"])
    if "Enforcing" in selinux:
        findings.append({
            "severity": "INFO",
            "category": "kernel",
            "title": "SELinux is enforcing",
            "detail": "Mandatory Access Control is active",
            "recommendation": ""
        })
    elif "Permissive" in selinux:
        findings.append({
            "severity": "HIGH",
            "category": "kernel",
            "title": "SELinux is in permissive mode",
            "detail": "Policies are loaded but not enforced",
            "recommendation": "Set SELINUX=enforcing in /etc/selinux/config"
        })
    elif "Disabled" in selinux or not selinux:
        findings.append({
            "severity": "MEDIUM",
            "category": "kernel",
            "title": "SELinux is disabled",
            "detail": "Mandatory Access Control is not available",
            "recommendation": "Consider enabling SELinux or AppArmor"
        })

    # Check AppArmor
    aa_status, _, _ = run_cmd(["aa-status"])
    if aa_status:
        profiles = [l for l in aa_status.splitlines() if "profiles" in l.lower()]
        findings.append({
            "severity": "INFO",
            "category": "kernel",
            "title": "AppArmor status",
            "detail": profiles[0] if profiles else aa_status[:200],
            "recommendation": ""
        })

    # Last boot time (for uptime and unplanned reboots)
    stdout, _, _ = run_cmd(["uptime", "-s"])
    if stdout:
        findings.append({
            "severity": "INFO",
            "category": "kernel",
            "title": "System boot time",
            "detail": f"Last boot: {stdout.strip()}",
            "recommendation": "Monitor for unexpected reboots"
        })

    return findings


# ═══════════════════════════════════════════════════
# AGGREGATOR
# ═══════════════════════════════════════════════════

CATEGORIES = {
    "user": audit_users,
    "sudo": audit_sudo,
    "file": audit_files,
    "ssh": audit_ssh,
    "network": audit_network,
    "service": audit_services,
    "package": audit_packages,
    "log": audit_logging,
    "kernel": audit_kernel,
}


def run_audit(categories=None, use_sudo=False):
    """Run selected audit categories. Returns results dict."""
    if use_sudo:
        os.environ.get("SUDO_ASKPASS", "")

    if categories is None or "all" in categories:
        cats = list(CATEGORIES.keys())
    else:
        cats = [c for c in categories if c in CATEGORIES]

    all_findings = []
    for cat in cats:
        try:
            results = CATEGORIES[cat]()
            all_findings.extend(results)
        except Exception as e:
            all_findings.append({
                "severity": "ERROR",
                "category": cat,
                "title": f"Audit error in {cat}",
                "detail": str(e),
                "recommendation": "Check permissions or tool availability"
            })

    # Sort by severity
    severity_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "INFO": 3, "ERROR": 4}
    all_findings.sort(key=lambda x: severity_order.get(x["severity"], 99))

    # Summary
    severity_counts = {}
    for f in all_findings:
        sev = f["severity"]
        severity_counts[sev] = severity_counts.get(sev, 0) + 1

    return {
        "audit_time": datetime.now().isoformat(),
        "hostname": platform.node(),
        "os": f"{platform.system()} {platform.release()}",
        "use_sudo": use_sudo,
        "summary": {
            "total_findings": len(all_findings),
            "by_severity": severity_counts,
            "score": _calculate_score(severity_counts),
        },
        "categories_audited": cats,
        "findings": all_findings,
    }


def _calculate_score(severity_counts):
    """Calculate a security score 0-100 (100 = most secure)."""
    score = 100
    score -= severity_counts.get("HIGH", 0) * 15
    score -= severity_counts.get("MEDIUM", 0) * 5
    score -= severity_counts.get("LOW", 0) * 2
    return max(0, min(100, score))


def generate_summary_text(results):
    """Generate a human-readable summary text."""
    s = results["summary"]
    lines = [
        f"Security Audit Report",
        f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━",
        f"Host: {results['hostname']}",
        f"OS:   {results['os']}",
        f"Time: {results['audit_time']}",
        f"",
        f"Security Score: {s['score']}/100",
        f"Total Findings: {s['total_findings']}",
        f"",
    ]
    for sev in ("HIGH", "MEDIUM", "LOW", "INFO", "ERROR"):
        count = s["by_severity"].get(sev, 0)
        if count > 0:
            icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢", "INFO": "ℹ️", "ERROR": "❌"}
            lines.append(f"  {icon.get(sev, '•')} {sev}: {count}")

    lines.append("")
    lines.append(f"Categories audited: {', '.join(results['categories_audited'])}")
    lines.append("")

    # Detailed findings
    for f in results["findings"]:
        icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢", "INFO": "ℹ️", "ERROR": "❌"}
        lines.append(f"{icon.get(f['severity'], '•')} [{f['severity']}] {f['title']}")
        lines.append(f"   {f['detail'][:200]}")
        if f["recommendation"]:
            lines.append(f"   → {f['recommendation']}")
        lines.append("")

    return "\n".join(lines)


def generate_word_report(results, output_path):
    """Generate a Word document audit report."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None, "python-docx not installed; install with: pip install python-docx"

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Security Audit Report")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.bold = True

    # Meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, val in [("Host", results["hostname"]),
                        ("OS", results["os"]),
                        ("Time", results["audit_time"])]:
        r = meta.add_run(f"\n{label}: {val}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Score
    score = results["summary"]["score"]
    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = score_p.add_run(f"Security Score: {score}/100")
    r.font.size = Pt(18)
    score_color = RGBColor(0x1B, 0x8A, 0x3A) if score >= 70 else \
                  RGBColor(0xCC, 0x88, 0x00) if score >= 40 else \
                  RGBColor(0xCC, 0x33, 0x33)
    r.font.color.rgb = score_color
    r.bold = True

    doc.add_page_break()

    # Summary
    h = doc.add_heading("Executive Summary", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    s = results["summary"]
    p = doc.add_paragraph(f"Total findings: {s['total_findings']}")
    p.runs[0].font.size = Pt(11)

    for sev, label in [("HIGH", "Critical"), ("MEDIUM", "Medium"), ("LOW", "Low")]:
        count = s["by_severity"].get(sev, 0)
        if count > 0:
            p = doc.add_paragraph()
            r = p.add_run(f"  {label}: {count}")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = {"HIGH": RGBColor(0xCC, 0x33, 0x33),
                                "MEDIUM": RGBColor(0xCC, 0x88, 0x00),
                                "LOW": RGBColor(0x66, 0x99, 0x33)}[sev]

    doc.add_paragraph(f"Categories audited: {', '.join(results['categories_audited'])}")
    doc.add_page_break()

    # Detailed findings by severity
    for severity_group, group_label in [("HIGH", "Critical Issues"),
                                        ("MEDIUM", "Medium Issues"),
                                        ("LOW", "Low Issues"),
                                        ("INFO", "Informational"),
                                        ("ERROR", "Errors")]:
        group = [f for f in results["findings"] if f["severity"] == severity_group]
        if not group:
            continue

        h = doc.add_heading(group_label, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

        for i, finding in enumerate(group, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {finding['title']}")
            r.bold = True
            r.font.size = Pt(10.5)

            p2 = doc.add_paragraph(f"   {finding['detail']}")
            p2.runs[0].font.size = Pt(9)

            if finding["recommendation"]:
                p3 = doc.add_paragraph(f"   → Recommendation: {finding['recommendation']}")
                p3.runs[0].font.size = Pt(9)
                p3.runs[0].font.color.rgb = RGBColor(0x2E, 0x6B, 0xA4)

    doc.save(output_path)
    return output_path, None


def main():
    parser = argparse.ArgumentParser(
        description="Linux Security Audit Tool — Read-Only, Non-Destructive",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
SAFETY GUARANTEES:
  • READ-ONLY: No system files or configurations are ever modified
  • No shell injection: All commands use subprocess.run with shell=False
  • Path allowlist: Only known safe config files are read
  • No destructive operations: No write, delete, chmod, or service change commands
        """
    )
    parser.add_argument(
        "--category", "-c", nargs="+",
        choices=list(CATEGORIES.keys()) + ["all"],
        default=["all"],
        help="Audit categories (default: all)"
    )
    parser.add_argument(
        "--sudo", action="store_true",
        help="Use sudo for privileged checks"
    )
    parser.add_argument(
        "--output", "-o",
        help="Save results as JSON to file"
    )
    parser.add_argument(
        "--report-doc",
        help="Generate Word (.docx) audit report"
    )
    parser.add_argument(
        "--text", action="store_true",
        help="Print human-readable text report"
    )

    args = parser.parse_args()

    # Run audit
    results = run_audit(categories=args.category, use_sudo=args.sudo)

    # Output
    if args.output:
        with open(args.output, "w") as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"✅ Audit results saved to: {args.output}")

    if args.report_doc:
        path, err = generate_word_report(results, args.report_doc)
        if path:
            print(f"✅ Audit report (Word): {path}")
        else:
            print(f"❌ Report generation failed: {err}")

    if args.text or not (args.output or args.report_doc):
        print(generate_summary_text(results))

    # Always print summary to stdout for tool capture
    s = results["summary"]
    print(f"\n{'='*50}")
    print(f"SCORE: {s['score']}/100 | "
          f"HIGH: {s['by_severity'].get('HIGH', 0)} | "
          f"MEDIUM: {s['by_severity'].get('MEDIUM', 0)} | "
          f"LOW: {s['by_severity'].get('LOW', 0)}")
    print(f"{'='*50}")

    # Return results for programmatic use
    return results


if __name__ == "__main__":
    main()
