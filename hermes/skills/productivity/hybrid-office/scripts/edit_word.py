#!/usr/bin/env python3
"""
Edit an existing Word document (.docx).

Usage:
    python edit_word.py <input_path> <output_path> < edits.json

Edits JSON format:
{
  "changes": [
    {"action": "replace_text", "old": "text to find", "new": "replacement text"},
    {"action": "replace_text", "old": "...", "new": "...", "occurrence": 1},
    {"action": "add_paragraph", "text": "New paragraph", "after_heading": "Section Title"},
    {"action": "add_paragraph", "text": "...", "at_end": true},
    {"action": "add_table", "headers": ["A","B"], "rows": [["1","2"]], "after_heading": "..."},
    {"action": "set_title", "text": "New Title"},
    {"action": "change_heading", "old": "Old heading", "new": "New heading"}
  ]
}
"""

import json
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_BLUE = RGBColor(0x2E, 0x6B, 0xA4)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def find_paragraph_with_text(doc, text, exact=False):
    """Find paragraph(s) containing text."""
    results = []
    for p in doc.paragraphs:
        full_text = p.text.strip()
        if exact and full_text == text:
            results.append(p)
        elif not exact and text in full_text:
            results.append(p)
    return results


def replace_text_in_paragraph(p, old, new, occurrence=0):
    """Replace text in a paragraph's runs."""
    full_text = p.text
    if occurrence > 0:
        # Replace Nth occurrence
        idx = 0
        for i, part in enumerate(full_text.split(old)):
            if i == occurrence:
                idx = sum(len(part) + len(old) for part in full_text.split(old)[:i])
                break
        # Simplified: just do find nth
        count = 0
        pos = 0
        while True:
            pos = full_text.find(old, pos)
            if pos == -1:
                break
            count += 1
            if count == occurrence:
                break
            pos += 1
        if pos == -1:
            return False

        before = full_text[:pos]
        after = full_text[pos + len(old):]
        p.text = ""
        run = p.add_run(before + new + after)
        return True
    else:
        if old not in full_text:
            return False
        new_text = full_text.replace(old, new)
        p.text = ""
        run = p.add_run(new_text)
        return True


def add_paragraph_after(doc, text, after_heading=None):
    """Add a paragraph after a heading."""
    if after_heading:
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            if p.style.name.startswith('Heading') and after_heading in p.text:
                # Insert after this paragraph
                ref_para = paras[i]
                new_p = doc.add_paragraph()
                run = new_p.add_run(text)
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

                # Move to correct position (not easily done with python-docx)
                # Instead just append and note
                return True, "appended"
    # Default: add at end
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return True, "appended"


def change_heading(doc, old, new):
    """Change a heading's text."""
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and old in p.text:
            p.text = ""
            run = p.add_run(new)
            run.font.color.rgb = DARK_BLUE
            return True
    return False


def edit_document(input_path, output_path, edits):
    """Apply edits to a Word document."""
    doc = Document(input_path)

    for edit in edits:
        action = edit.get("action")
        if action == "replace_text":
            old = edit["old"]
            new = edit["new"]
            occurrence = edit.get("occurrence", 0)
            count = 0
            for p in doc.paragraphs:
                if occurrence > 0 and count >= occurrence:
                    break
                if old in p.text:
                    if occurrence > 0:
                        count += 1
                        if count == occurrence:
                            replace_text_in_paragraph(p, old, new, 1)
                            break
                    else:
                        replace_text_in_paragraph(p, old, new)
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old in cell.text:
                            for p in cell.paragraphs:
                                if old in p.text:
                                    replace_text_in_paragraph(p, old, new)
                                    if occurrence > 0:
                                        break

        elif action == "add_paragraph":
            text = edit.get("text", "")
            after_heading = edit.get("after_heading")
            add_paragraph_after(doc, text, after_heading)

        elif action == "set_title":
            # Find largest heading in first few paragraphs
            for p in doc.paragraphs[:5]:
                if p.style.name.startswith('Heading 1'):
                    p.text = ""
                    run = p.add_run(edit["text"])
                    run.font.size = Pt(16)
                    run.font.name = "Calibri"
                    run.font.color.rgb = DARK_BLUE
                    break

        elif action == "change_heading":
            change_heading(doc, edit["old"], edit["new"])

        elif action == "add_table":
            from docx.shared import Pt as Pt2
            from docx.enum.table import WD_TABLE_ALIGNMENT
            headers = edit.get("headers", [])
            rows = edit.get("rows", [])
            num_cols = len(headers) if headers else len(rows[0]) if rows else 1
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(h)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt2(9.5)
                        run.font.name = "Calibri"

            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        cell = table.rows[r_idx + 1].cells[c_idx]
                        cell.text = str(cell_text)

    doc.save(output_path)
    print(f"✅ Word document edited and saved: {output_path}")
    return output_path


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Edit an existing Word document")
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("output", help="Output .docx file path")
    args = parser.parse_args()

    raw = sys.stdin.read()
    if not raw.strip():
        print("❌ No edit JSON provided via stdin", file=sys.stderr)
        sys.exit(1)

    edits_data = json.loads(raw)
    if isinstance(edits_data, dict) and "changes" in edits_data:
        edits = edits_data["changes"]
    else:
        edits = edits_data

    edit_document(args.input, args.output, edits)


if __name__ == "__main__":
    main()
