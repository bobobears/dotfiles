#!/usr/bin/env python3
"""
Create a professionally formatted Word document (.docx).

Usage:
    python create_word.py <output_path> [--title "Title"] [--no-page-numbers]

Reads document structure from stdin as JSON:
{
  "title": "Document Title",
  "subtitle": "Optional Subtitle",
  "author": "Author Name",
  "sections": [
    {
      "heading": "Section Title",
      "level": 1,
      "paragraphs": [
        {"text": "Normal paragraph text.", "style": "normal"},
        {"text": "Bold lead-in text", "bold": true, "text_after": " then continues normal"},
        {"text": "Bullet point", "style": "bullet"},
        {"text": "Numbered item", "style": "numbered"}
      ],
      "tables": [
        {
          "headers": ["Col1", "Col2", "Col3"],
          "rows": [
            ["Data1", "Data2", "Data3"]
          ],
          "caption": "Table caption",
          "style": "professional"
        }
      ]
    }
  ]
}
"""

import json
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Professional Color Palette ──
COLORS = {
    "dark_blue":    RGBColor(0x1B, 0x3A, 0x5C),  # Headings
    "accent_blue":  RGBColor(0x2E, 0x6B, 0xA4),  # Table header bg
    "light_bg":     RGBColor(0xF2, 0xF6, 0xFA),  # Table alt row
    "body_text":    RGBColor(0x33, 0x33, 0x33),  # Body text
    "muted":        RGBColor(0x66, 0x66, 0x66),  # Captions, footers
    "white":        RGBColor(0xFF, 0xFF, 0xFF),
    "border":       RGBColor(0xBB, 0xBB, 0xBB),
}


def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right (dicts with sz, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{attrs.get("sz", "4")}" '
            f'w:space="0" w:color="{attrs.get("color", "BBBBBB")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def apply_table_style(table, style_name="professional"):
    """Apply professional styling to a table."""
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            # Clear default paragraph
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = "Calibri"

            if row_idx == 0:
                # Header row
                set_cell_shading(cell, "2E6BA4")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.color.rgb = COLORS["white"]
                        run.font.bold = True
                        run.font.size = Pt(9.5)
                        run.font.name = "Calibri"
                # Bottom border on header
                set_cell_border(cell, bottom={"sz": "8", "color": "1B3A5C"})
            else:
                # Alternating row colors
                if row_idx % 2 == 0:
                    set_cell_shading(cell, "F2F6FA")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = COLORS["body_text"]
                        run.font.size = Pt(9.5)
                        run.font.name = "Calibri"


def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = COLORS["muted"]
            run.font.name = "Calibri"


def add_paragraph_with_format(doc, text, bold=False, italic=False,
                               font_size=Pt(10.5), font_name="Calibri",
                               color=COLORS["body_text"],
                               alignment=WD_ALIGN_PARAGRAPH.LEFT,
                               space_after=Pt(6), space_before=Pt(0),
                               first_line_indent=None):
    """Add a paragraph with full formatting control."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.15

    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent

    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with professional styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = COLORS["dark_blue"]
        if level == 1:
            run.font.size = Pt(16)
            run.font.name = "Calibri"
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.name = "Calibri"
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
    return h


def build_document(structure, output_path, show_page_numbers=True):
    """Build a complete Word document from a structure dict."""
    doc = Document()

    # ── Default style tuning ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ── Title page ──
    title = structure.get("title", "Document")
    add_paragraph_with_format(doc, "", font_size=Pt(72), space_before=Pt(72))
    add_paragraph_with_format(doc, title, bold=True, font_size=Pt(26),
                              font_name="Calibri", color=COLORS["dark_blue"],
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_after=Pt(8))

    subtitle = structure.get("subtitle")
    if subtitle:
        add_paragraph_with_format(doc, subtitle, font_size=Pt(14),
                                  font_name="Calibri", color=COLORS["accent_blue"],
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_after=Pt(18))

    author = structure.get("author")
    if author:
        add_paragraph_with_format(doc, author, font_size=Pt(10),
                                  font_name="Calibri", color=COLORS["muted"],
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_after=Pt(6))

    # Page break after title
    doc.add_page_break()

    # ── Sections ──
    for section in structure.get("sections", []):
        add_heading_styled(doc, section.get("heading", ""),
                           level=section.get("level", 1))

        for para in section.get("paragraphs", []):
            style_type = para.get("style", "normal")

            if style_type == "bullet":
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(para.get("text", ""))
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"
                run.font.color.rgb = COLORS["body_text"]

            elif style_type == "numbered":
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(para.get("text", ""))
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"
                run.font.color.rgb = COLORS["body_text"]

            elif style_type == "quote":
                add_paragraph_with_format(doc, para.get("text", ""),
                                          italic=True, font_size=Pt(10),
                                          color=COLORS["muted"],
                                          alignment=WD_ALIGN_PARAGRAPH.LEFT)

            else:
                text = para.get("text", "")
                if para.get("bold"):
                    # Bold lead-in
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    r1 = p.add_run(para.get("text", ""))
                    r1.bold = True
                    r1.font.size = Pt(10.5)
                    r1.font.name = "Calibri"
                    r1.font.color.rgb = COLORS["body_text"]
                    if para.get("text_after"):
                        r2 = p.add_run(para["text_after"])
                        r2.font.size = Pt(10.5)
                        r2.font.name = "Calibri"
                        r2.font.color.rgb = COLORS["body_text"]
                else:
                    add_paragraph_with_format(doc, text)

        # ── Tables ──
        for tbl in section.get("tables", []):
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if not headers and not rows:
                continue

            caption = tbl.get("caption")
            if caption:
                add_paragraph_with_format(doc, caption, italic=True,
                                          font_size=Pt(9), color=COLORS["muted"],
                                          space_before=Pt(10), space_after=Pt(4))

            num_cols = len(headers) if headers else len(rows[0]) if rows else 1
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            # Auto-fit
            table.autofit = True

            # Header row
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(h))
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"
                run.font.color.rgb = COLORS["white"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, "2E6BA4")

            # Data rows
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        cell = table.rows[r_idx + 1].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(str(cell_text))
                        run.font.size = Pt(9.5)
                        run.font.name = "Calibri"
                        run.font.color.rgb = COLORS["body_text"]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        # Alternating row color
                        if r_idx % 2 == 0:
                            set_cell_shading(cell, "F2F6FA")

            # Space after table
            add_paragraph_with_format(doc, "", font_size=Pt(4), space_after=Pt(2))

    if show_page_numbers:
        add_page_number(doc)

    doc.save(output_path)
    print(f"✅ Word document saved: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Create a professional Word document")
    parser.add_argument("output", help="Output .docx file path")
    parser.add_argument("--title", help="Document title (override JSON)")
    parser.add_argument("--no-page-numbers", action="store_true",
                        help="Disable page numbers")
    args = parser.parse_args()

    # Read structure from stdin
    raw = sys.stdin.read()
    if not raw.strip():
        print("❌ No input JSON provided via stdin", file=sys.stderr)
        sys.exit(1)

    structure = json.loads(raw)
    if args.title:
        structure["title"] = args.title

    build_document(structure, args.output,
                   show_page_numbers=not args.no_page_numbers)


if __name__ == "__main__":
    main()
